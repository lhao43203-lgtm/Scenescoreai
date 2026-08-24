from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(r"E:\安装\Scenescoreai")
OUTPUT = Path(
    r"C:\Users\Administrator\Documents\xwechat_files\wxid_ugaub4ep4d8l22_7f65\msg\file\2026-08\Scene_Score_Web_部署PRD_v1.3_交互更新.docx"
)

INK = "111111"
MUTED = "666666"
RED = "F44336"
PAPER = "F5F5F2"
LINE = "D9D9D3"
BLUE = "315C72"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=LINE, size="6"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_font(run, size=10.5, bold=False, color=INK, italic=False):
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def style_paragraph(paragraph, before=0, after=6, line=1.35):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_text(doc, text, size=10.5, bold=False, color=INK, italic=False, align=None, after=6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    style_paragraph(p, after=after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color, italic=italic)
    return p


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    style_paragraph(p, after=3, line=1.3)
    set_run_font(p.add_run(text), size=10.5)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    style_paragraph(p, before=12 if level == 1 else 8, after=7, line=1.1)
    run = p.add_run(text)
    set_run_font(run, size=17 if level == 1 else 13, bold=True, color=RED if level == 1 else INK)
    return p


def add_note(doc, text):
    p = doc.add_paragraph()
    style_paragraph(p, before=3, after=9, line=1.35)
    p.paragraph_format.left_indent = Cm(0.35)
    run = p.add_run(text)
    set_run_font(run, size=9.5, color=MUTED, italic=True)
    return p


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for i, value in enumerate(headers):
        header_cells[i].text = ""
        set_cell_shading(header_cells[i], RED)
        set_cell_border(header_cells[i], RED)
        header_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = header_cells[i].paragraphs[0]
        style_paragraph(p, after=2, line=1.15)
        set_run_font(p.add_run(value), size=9.5, bold=True, color="FFFFFF")
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            set_cell_shading(cells[i], "FFFFFF" if row_index % 2 == 0 else PAPER)
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            p = cells[i].paragraphs[0]
            style_paragraph(p, after=2, line=1.25)
            set_run_font(p.add_run(str(value)), size=9.2)
    if widths:
        for row in table.rows:
            for cell, width in zip(row.cells, widths):
                cell.width = Cm(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_section_label(doc, text):
    p = doc.add_paragraph()
    style_paragraph(p, before=14, after=3, line=1)
    run = p.add_run(text.upper())
    set_run_font(run, size=8.5, bold=True, color=BLUE)
    return p


def add_page_break(doc):
    doc.add_page_break()


def add_ranking_section(doc):
    """Section 06 is copied from v1.2 without product/content changes."""
    add_heading(doc, "06 | 四个独立榜单", 1)
    add_table(
        doc,
        ["榜单入口", "路由", "默认周期", "展示内容"],
        [
            ["AI漫剧榜", "/ranking/anime", "Daily", "排名、作品、类型、年份、Scene Score、周期"],
            ["AI短剧榜", "/ranking/short", "Daily", "排名、作品、类型、年份、Scene Score、周期"],
            ["AI音乐榜", "/ranking/music", "Daily", "排名、作品、类型、年份、Scene Score、周期"],
            ["生产工具榜", "/ranking/tools", "Daily", "排名、作品、类型、年份、Scene Score、周期"],
        ],
    )
    add_table(
        doc,
        ["period_type", "period_key", "展示范围"],
        [
            ["daily", "YYYY-MM-DD", "指定自然日榜单"],
            ["monthly", "YYYY-MM", "指定自然月榜单"],
            ["yearly", "YYYY", "指定自然年榜单"],
        ],
    )
    add_table(
        doc,
        ["字段", "要求"],
        [
            ["Rank", "显示当前周期中的排名，使用已发布数据的 rank。"],
            ["Title", "显示作品名称；点击进入 Series Detail。"],
            ["Poster / Media", "显示作品媒体；加载失败时显示明确占位内容。"],
            ["Creator", "显示导演、创作者或生产方（字段可用时展示）。"],
            ["Type / Year", "显示作品类型和发行年份。"],
            ["Scene Score", "显示十分制最终分数，保留 1 位小数，例如 8.7。"],
            ["Period", "显示当前选中的日榜、月榜或年榜。"],
            ["Advertisement", "广告不占用作品 rank 编号，并显示 Ad 标识。"],
        ],
    )
    add_table(
        doc,
        ["功能", "规则"],
        [
            ["默认周期", "首次进入任一榜单默认 Daily。"],
            ["切换周期", "切换后读取对应已发布数据；无数据时显示 Empty 状态。"],
            ["筛选", "支持按榜单已有字段进行筛选；筛选不重新计算 rank 和 score。"],
            ["分页", "列表内容较多时使用分页或 Load More，并保留当前榜单类型和周期。"],
            ["作品入口", "点击卡片或查看详情进入 /series/:id；返回时保留可恢复的列表状态。"],
        ],
    )
    add_table(
        doc,
        ["编号", "验收条件"],
        [
            ["R-01", "四个独立榜单入口均可访问。"],
            ["R-02", "日榜、月榜、年榜均可切换，并读取正确的 period_type 和 period_key。"],
            ["R-03", "列表显示 rank、title、score，并可进入作品详情。"],
            ["R-04", "广告不改变作品排名和排名编号。"],
            ["R-05", "无数据、接口错误、图片失败时显示明确状态。"],
        ],
    )


def build_document():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(INK)

    # Cover
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    style_paragraph(p, after=50)
    set_run_font(p.add_run("SCENE SCORE  /  WEB PRD"), size=10, bold=True, color=RED)
    add_text(doc, "正式网页产品需求文档", size=30, bold=True, after=4)
    add_text(doc, "Scene Score Web Product v1.3", size=20, bold=True, after=3)
    add_text(doc, "交互与部署更新版  ·  公开访客型 Web 网站", size=11, color=MUTED, after=32)
    add_note(doc, "本版本基于当前 Scene Score Vite + React 前端的真实页面、路由、交互和媒体结构更新。原 v1.2 文件保留不改；第 06 节“四个独立榜单”及其字段、规则、验收内容保持原文不变。")
    add_table(
        doc,
        ["项目", "内容"],
        [
            ["产品名称", "Scene Score"],
            ["文档版本", "v1.3 / 交互更新版"],
            ["产品形态", "公开访客型 Web 网站；PC 优先，兼容 Tablet 与 Mobile"],
            ["前端底座", "Vite + React + TypeScript；不迁移到 Next.js"],
            ["文档用途", "产品、前端、后端、设计、内容运营、部署和测试对齐"],
            ["更新时间", "2026-08-20"],
        ],
    )
    add_page_break(doc)

    # 01 Product scope
    add_heading(doc, "01 | 产品定位与本次修订范围", 1)
    add_text(doc, "Scene Score 是面向屏幕故事、影像作品和创作人员的公开索引体验。网站以沉浸式首页、作品索引、评审资料和可解释的评分方法建立品牌入口；公开端不要求访客注册或登录，暂不展示未确认的真实评分。", after=8)
    add_heading(doc, "1.1 本次修订目标", 2)
    for item in [
        "以当前网页实际实现为准，补齐首页滚动、页面转场、菜单状态、语言切换、媒体降级和详情跳转说明。",
        "把作品列表首页明确为 /explore，说明网格 / 列表切换、筛选、Hover 视频预览和作品详情闭环。",
        "把评委列表与评委详情拆成两个明确页面；评委顺序按当前内容资料，Paco Wong 排第一。",
        "评分方法页改为公开框架预览，不虚构正式权重、分数或评审结果；正式榜单规则仍以第 06 节为准。",
        "覆盖英文 / 繁中全局切换，并要求繁中标题、按钮、段落和导航使用独立的字距与行高。",
        "补充可部署运行方式、SPA 回退、静态媒体、视频 Range 请求和生产验收标准。",
    ]:
        add_bullet(doc, item)
    add_heading(doc, "1.2 本期不包含", 2)
    for item in [
        "公开访客注册、登录、找回密码、个人中心和用户账号体系。",
        "用户投稿、用户评分、评论、点赞、收藏、关注、私信等社区功能。",
        "外部影视数据库、自动化评分服务和前台实时计算排名。",
        "运营后台、权限系统、真实评分数据和内部 Draft / Hidden 内容的公开暴露。",
        "第 06 节四个排行榜的业务规则改写；该章节完整沿用 v1.2。",
    ]:
        add_bullet(doc, item)

    # 02 roles
    add_heading(doc, "02 | 用户角色与访问边界", 1)
    add_table(
        doc,
        ["角色", "身份", "权限"],
        [
            ["公开访客", "无需登录的网站访问者", "访问首页、作品索引、作品详情、评委、评分方法和关于页面；可切换英文 / 繁中；可播放或预览公开媒体。"],
            ["产品 / 评审人员", "使用本地或部署预览站查看 UI 与流程", "检查页面跳转、Hover / Click 状态、滚动动画、媒体降级和繁中排版；不修改公开数据。"],
            ["运营人员", "内部内容维护人员", "后续接入后维护作品、媒体、评委、评分方法、站点文案和发布状态。"],
            ["后端工程师", "提供公开 API 与静态资源服务", "保证数据字段、Published 边界、媒体 URL、SPA fallback、Range 请求和 SEO 元信息。"],
        ],
    )
    add_note(doc, "公开页面只允许读取 Published 内容。Draft、Hidden、Expired 或未确认的内部字段，不得通过页面、接口或静态资源 URL 暴露。")

    # 03 IA and routes
    add_heading(doc, "03 | 信息架构与当前正式路由", 1)
    add_text(doc, "当前顶部导航使用五个公开入口；旧版 /contact、/rules 和 /ranking/* 路径保留兼容跳转，避免外部链接失效。排行榜业务定义仍见第 06 节。", after=7)
    add_table(
        doc,
        ["导航 / 页面", "当前路由", "页面目的", "主要入口"],
        [
            ["首页 Home", "/", "品牌入口、沉浸式视频开场、作品摘要、评分方法摘要、评审入口和 Footer CTA", "Logo、首页导航、转场完成后回到首页"],
            ["作品索引 Explore", "/explore", "当前列表首页；作品网格 / 列表、筛选、媒体预览和作品详情入口", "导航、首页作品卡片、查看完整排名"],
            ["作品详情 Series Detail", "/series/:id", "单项作品媒体、概览、制作团队抽屉、静态画面和下一项入口", "Explore 卡片、首页作品卡片、详情页 Next"],
            ["特邀评审 Judges", "/judges", "评委名单、Hover / Focus 切换、选中评委图片和简介", "主导航、首页认识评审"],
            ["评委详情 Judge Detail", "/judges/:id", "单位评委的完整长页面资料、段落、前后评委导航", "评委左侧名单、评委预览图片、View Full Profile"],
            ["评分方法 Methodology", "/methodology", "公开方法框架预览与状态边界", "主导航、首页评分方法 CTA、作品详情 Behind the Score"],
            ["关于我们 About", "/about", "品牌定位、价值观、合作入口和联系锚点", "主导航、Footer CTA"],
            ["兼容入口", "/contact", "重定向至 /about#contact；不再单独维护旧版联系页面", "旧链接和外部书签"],
            ["兼容入口", "/rules", "重定向至 /methodology", "旧规则链接"],
            ["兼容入口", "/ranking/*", "重定向至 /explore；排行榜业务章节保留 v1.2 定义", "旧榜单链接"],
        ],
    )
    add_table(
        doc,
        ["公共路由场景", "系统行为"],
        [
            ["直接访问页面", "无需账号校验；加载对应页面；生产服务器必须将未知前端路径回退至 index.html。"],
            ["不存在的作品或评委 ID", "返回对应列表页，不渲染空详情；不得泄露内部数据。"],
            ["点击导航", "进入目标路由，触发 1.25 秒纸张转场；转场期间锁定滚动输入，完成后解锁。"],
            ["提前滚动 / 滚轮输入", "路由切换开始时只初始化一次滚动位置；转场完成事件不会再次重置滚动，避免跳回首屏。"],
            ["语言切换", "写入 localStorage 的 app-language；全局更新页面、导航、按钮、评委内容和 Footer。Logo 与 Scene Score 品牌名保持英文。"],
        ],
    )

    # 04 flows
    add_heading(doc, "04 | 当前用户流程与交互验收", 1)
    add_table(
        doc,
        ["流程编号", "用户动作", "系统结果"],
        [
            ["F-01", "首次进入 /", "加载首页沉浸式 Hero；视频默认静音、循环、尝试自动播放；视频尚未就绪时显示静态色彩 / 封面降级，不出现空白。"],
            ["F-02", "向下滚动首页首屏", "第一阶段 pin 住视频，只移动橙色遮罩板；遮罩从右下方向左上退出，并带走遮罩内的 Scene Score、REC 和辅助文字。"],
            ["F-03", "继续向下滚动", "橙色遮罩完全消失后才释放视频滚动阶段；视频不因遮罩滚动发生抖动、反复放大或布局跳变。"],
            ["F-04", "点击菜单切换页面", "新页面像纸张一样从左下方进入，旧页面向右退出；过程约 1.25 秒，页面切换时保持矩形 active 框和红点状态。"],
            ["F-05", "从其他页面回到 Home", "首页视频保留上次 currentTime，转场结束后继续播放，不从头开始；首页 ScrollTrigger 等待 route-transition-complete 后再初始化。"],
            ["F-06", "Hover / Focus 导航和作品卡片", "导航显示对应的 Hover 状态；作品卡片使用稳定图片首帧，鼠标进入时播放对应短视频，鼠标离开时暂停，不阻止卡片点击。"],
            ["F-07", "进入 /explore", "默认展示 All 作品；可切换 Featured、Short Film、Series，支持 RESET、TOP、GRID、LIST；作品卡片点击进入 /series/:id。"],
            ["F-08", "进入作品详情", "显示作品媒体、返回索引、作品概览、Behind the Score、制作团队抽屉、静态画面和下一项；缺失媒体时继续显示 poster 或占位。"],
            ["F-09", "进入 /judges", "默认选中 Paco Wong；Hover / Focus 左侧名单时切换右侧评委画面与简介；点击左侧名称或右侧图片进入对应 /judges/:id。"],
            ["F-10", "进入评委详情", "显示一页连续的完整介绍内容；支持返回评审总览、上一位 / 下一位评审和查看评审总览。英文与繁中内容保持同等详细程度。"],
            ["F-11", "进入 /methodology", "显示四个公开框架卡片：评分维度、权重逻辑、评审规程、异常与冲突；明确这是公开预览，不虚构最终评分。"],
            ["F-12", "切换英文 / 繁中", "页面立即更新；繁中标题取消英文负字距，采用更宽的 letter-spacing、自然行高和 keep-all 换行，不能出现笔画黏连或裁切。"],
        ],
    )
    add_note(doc, "页面转场的技术术语：route transition / page transition；首页遮罩退出属于 scroll-linked animation，视频固定阶段属于 pinned media stage。实现上优先使用 transform、opacity 和 GSAP ScrollTrigger，避免在滚动中改变 top / left / width / height。")

    # 05 Home
    add_heading(doc, "05 | Home 首页（非排行榜部分更新）", 1)
    add_table(
        doc,
        ["区块", "当前内容", "交互与状态"],
        [
            ["Header", "Logo、首页、探索、特邀评审、评分方法、关于我们；英文 / 繁中切换。", "当前页使用红点和矩形边框；Hover 显示对应菜单状态；移动端收起为可展开导航。"],
            ["Immersive Hero", "首页背景视频、橙色遮罩板、Scene Score 标题、REC 状态和辅助文字。", "自动播放失败时保留静态背景；遮罩先滚动离场，视频保持固定；遮罩消失后释放后续页面滚动。"],
            ["首页作品摘要", "Highlight、前五项作品、作品媒体和查看入口。", "卡片与文字链接均可进入 /explore 或 /series/:id；图片是静态首帧，视频用于 Hover 预览。"],
            ["评分方法摘要", "透明始于设计 / 具体落实于实践；链接到 /methodology。", "不展示虚构分数；状态由公开框架预览文案说明。"],
            ["评审摘要", "五位评委图片堆叠、介绍和认识评审 CTA。", "图片与对应评委详情可点击；顺序与评委资料保持一致。"],
            ["Footer", "公开索引、站点链接、合作入口和版权信息。", "不再显示 LOCAL MVP / 本地 MVP；仅保留公开索引、公开浏览和联系入口。"],
        ],
    )
    add_table(
        doc,
        ["编号", "验收条件"],
        [
            ["H-01", "首页首次加载不出现大面积空白；视频失败时仍能看到遮罩、标题和静态背景。"],
            ["H-02", "滚动橙色遮罩时视频画面保持稳定；遮罩消失后才进入视频自然滚动。"],
            ["H-03", "从其他页面切回 Home，视频从离开时的位置继续，不明显卡顿或从头开始。"],
            ["H-04", "提前滚轮输入不会在转场完成后再次跳回首屏。"],
            ["H-05", "首页所有 CTA、作品图片、评审图片和导航均可点击或键盘访问。"],
        ],
    )

    # Ranking section unchanged
    add_ranking_section(doc)

    # 07 detail
    add_heading(doc, "07 | Series Detail 作品详情（按当前交互重写）", 1)
    add_table(
        doc,
        ["区块", "内容 / 字段", "交互规则"],
        [
            ["Hero", "作品视频、poster、年份、类型、编号、作品标题、REC / Preview Frame。", "视频默认静音循环；poster 作为加载和失败降级；返回按钮返回 /explore。"],
            ["Overview", "作品说明、透明度说明、Behind the Score。", "文字链接进入 /methodology；繁中版本保持完整句子，不使用英文压缩字距。"],
            ["Crew drawer", "制作团队标题、导演和制作资料占位。", "点击 THE CREW / 製作團隊 打开全屏抽屉；右上角关闭；不跳转路由。"],
            ["Stills", "当前作品和下一项作品的两张媒体首帧。", "采用 object-fit cover；图片失败时保留色彩背景和编号。"],
            ["Next asset", "下一项作品入口。", "点击进入下一个 /series/:id；保持作品索引顺序。"],
        ],
    )
    add_note(doc, "当前作品数据位于 src/data/sceneScore.ts；媒体首帧位于 public/images/works/，Hover 视频位于 public/media/ranking/。后续后端替换时只需保持 image 与 video 为稳定 HTTP/CDN URL。")

    # 08 Judges
    add_heading(doc, "08 | 特邀评审与评委详情", 1)
    add_table(
        doc,
        ["页面", "内容", "交互"],
        [
            ["/judges", "五位评委名单、右侧当前人物图片、职位、Headline 和简介。", "默认第一位为 Paco Wong；Hover / Focus 切换预览；点击名单或图片进入详情。"],
            ["/judges/:id", "完整评委介绍，包含当前人物、专业身份、多个资料段落和图片。", "保持长页面阅读；可返回总览、进入上一位或下一位评委；英文与繁中段落数量和信息量对齐。"],
            ["评委顺序", "01 Paco Wong；02 Edmond Wong；03 Bennett Pang；04 CK Chan；05 Chan Tai Lee。", "列表编号、首页头像堆叠、详情页上一位 / 下一位均使用同一顺序。"],
        ],
    )
    add_heading(doc, "8.1 评委媒体与排版要求", 2)
    for item in [
        "评委照片不使用黑白滤镜；保持彩色和原始人物辨识度。",
        "照片使用 object-fit: cover 并根据原图长宽比调整 object-position，避免人物头部和主体被裁切。",
        "繁中姓名、职位、Headline 和段落使用独立字距、行高和 keep-all；不将中文强制套用英文负字距。",
        "点击图片和点击左侧评委名称的行为保持一致，均进入对应详情页。",
    ]:
        add_bullet(doc, item)

    # 09 Methodology / About / i18n
    add_heading(doc, "09 | 评分方法、关于我们与全局语言", 1)
    add_table(
        doc,
        ["页面", "当前展示", "验收要求"],
        [
            ["/methodology", "公开评分框架预览：评分维度、权重逻辑、评审规程、异常与冲突。", "明确公开预览状态；正式协议未确认前不渲染虚构分数。"],
            ["/about", "品牌定位、价值观、项目说明和合作 CTA。", "联系 CTA 使用 mailto；删除旧的 MVP / 本地原型措辞。"],
            ["全局繁中", "导航、按钮、说明、评委内容、作品详情、Footer 和状态标签。", "切换后全局一致；Logo、Scene Score、品牌专有名词保留英文；标题不重叠、不被截断。"],
        ],
    )
    add_note(doc, "繁中 Typography 规范：中文展示标题建议 line-height 1.08–1.25、letter-spacing 0.045em–0.06em；正文 line-height 1.7–1.85；避免 text-transform: uppercase 和负字距影响中文。")

    # 10 media and states
    add_heading(doc, "10 | 媒体、加载和异常状态", 1)
    add_table(
        doc,
        ["状态", "触发条件", "前台表现"],
        [
            ["Default", "数据和媒体正常", "显示页面内容、图片首帧和可用交互。"],
            ["Loading", "路由、图片或视频首次加载", "保持容器比例，避免页面跳动；视频未就绪时使用图片或色彩背景。"],
            ["Hover preview", "鼠标进入 Explore 作品卡片", "播放对应短视频；鼠标离开暂停；不影响点击详情。"],
            ["Media failure", "图片 / 视频加载失败", "显示静态 poster、渐变背景、编号和文字信息，不出现空白区域。"],
            ["Not Found", "作品或评委 ID 不存在", "返回对应索引页；不泄露未发布字段。"],
            ["Reduced motion", "系统开启 prefers-reduced-motion", "减少纸张转场、滚动 reveal 和 REC 动画；保留内容和可用导航。"],
        ],
    )
    add_heading(doc, "10.1 当前素材路径", 2)
    for item in [
        "首页视频：public/media/scene-score-home.mp4。",
        "排行榜作品首帧：public/images/works/work-01.png 至 work-05.png。",
        "排行榜作品视频：public/media/ranking/work-01.mp4 至 work-05.mp4。",
        "评委图片：public/images/judges/；Logo：public/images/logo/。",
        "正式公网发布前必须再次确认所有图片、视频、字体和品牌素材的授权范围。",
    ]:
        add_bullet(doc, item)

    # 11 deployment
    add_heading(doc, "11 | 部署、环境与后端交接", 1)
    add_table(
        doc,
        ["项目", "要求"],
        [
            ["Node.js", "建议 Node.js 24；npm 随 Node.js 安装。"],
            ["安装", "项目根目录执行 npm install；不复制 node_modules、dist 或缓存。"],
            ["检查", "npm run check；包含 ESLint、TypeScript project build 和 Vite production build。"],
            ["构建", "npm run build 生成 dist/；npm run preview 用于生产静态预览。"],
            ["SPA fallback", "Nginx、IIS、静态托管必须将 /explore、/series/:id、/judges/:id 等未知路径回退到 index.html。"],
            ["媒体服务", "视频服务应支持正确 Content-Type 和 Range 请求；静态资源 URL 使用 HTTP/CDN 路径，不使用 Windows 本地路径。"],
            ["安全", "前端不存放密钥、账号或内部接口；外部 target_url 需做协议和域名校验。"],
            ["后端边界", "后端后续提供 Published 作品、媒体、评委、方法和站点内容；前端保持字段 image、video、status、id 等稳定。"],
        ],
    )
    add_text(doc, "本地演示：", size=10.5, bold=True, after=3)
    add_text(doc, "cd E:\\安装\\Scenescoreai\nnpm install\nnpm run dev -- --host 127.0.0.1\n浏览器访问 http://127.0.0.1:5173/", size=10, color=BLUE, after=8)
    add_note(doc, "前端交接文档：docs/TECH-HANDOFF.md；产品经理本地演示：docs/PRODUCT-DEMO.md；后端交接：docs/BACKEND-HANDOFF.md。")

    # 12 acceptance
    add_heading(doc, "12 | 新版交互验收标准", 1)
    add_table(
        doc,
        ["类别", "验收标准"],
        [
            ["首页滚动", "橙色遮罩先移动并完全退出；视频在第一阶段固定；遮罩消失后才进入视频滚动；无明显抖动。"],
            ["路由转场", "页面从左下进入、旧页面向右退出；时长约 1.25 秒；转场期间滚动锁定，结束后可操作。"],
            ["Home 视频", "从 Home 切出后暂停并保存 currentTime；切回后恢复原位置，不从头播放；不能因转场重复初始化造成卡顿。"],
            ["导航状态", "当前菜单有红点和矩形框；Hover / Focus 状态可见；桌面和移动端均可操作。"],
            ["作品索引", "/explore 可访问；筛选、RESET、TOP、GRID、LIST 可用；图片首帧显示；Hover 播放视频；点击进入详情。"],
            ["作品详情", "返回、制作团队抽屉、Behind the Score、静态画面和下一项入口可用；媒体失败有降级。"],
            ["评委流程", "/judges 初始 Paco Wong；Hover / Focus 切换；点击名单和图片均能进入对应完整详情页。"],
            ["语言", "英文 / 繁中切换覆盖全局；繁中标题、正文、菜单和按钮不黏连、不溢出、不被裁切。"],
            ["排行榜", "第 06 节原有四榜、周期、字段、规则和验收内容保持不变；旧 /ranking/* 路径可兼容跳转。"],
            ["部署", "npm run check 通过；生产构建可生成 dist；直接刷新公开路由不 404；图片和视频 URL 正常返回。"],
        ],
    )
    add_heading(doc, "13 | 版本变更记录", 1)
    add_table(
        doc,
        ["版本", "变更"],
        [
            ["v1.3", "依据当前 Scene Score 前端重写非排行榜章节：真实路由、首页遮罩滚动、纸张转场、视频续播、作品索引、作品详情、评委详情、评分方法、繁中和部署。"],
            ["v1.3", "第 06 节四个独立榜单保持 v1.2 原文，不调整其业务范围、字段、周期、规则和验收内容。"],
        ],
    )
    add_text(doc, "—— 文档结束 ——", size=10, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, after=0)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
